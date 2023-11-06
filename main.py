from pathlib import Path
from docx import Document
from collections import defaultdict

from src.class_moodle_question import MoodleQuestion
from src.export import export_to_xml

from conf import DATA_DIR, POINTS_IN_TOPIC

def get_list_file():
    res = []
    for f in Path(DATA_DIR).glob('*.docx'):
        res.append(f)
    return res

def clearing_text(text):
    def clearing(s:str):               
        i = 0
        while i <len(s) and not s[i].isalpha():
            i +=1
        return s[i:]
    if isinstance(text,list):
        for i in range(len(text)):
            text[i] = clearing(text[i]) 
        return text   
    else:
        return clearing(text)


def parse_doc(document,name_file,questions:list):
    i = 0
    while i < len(document.paragraphs):
        question = MoodleQuestion()
        while document.paragraphs[i].text == '':  
            i += 1  
            if i >= len(document.paragraphs): break      
        question.question = clearing_text(document.paragraphs[i].text)
        list_anwer_options = []
        i += 1
        if i >= len(document.paragraphs): break
        while document.paragraphs[i].text != '':      
            list_anwer_options.append(document.paragraphs[i].text)
            i += 1
            if i >= len(document.paragraphs): break
        question.answer_options = clearing_text(list_anwer_options[:-1])
        question.correct_answer_number = int(list_anwer_options[-1])
        question.tag = name_file
        questions.append(question)

def distribute_points(lq:list[MoodleQuestion]) -> None:
    
    
    topics = defaultdict(list)
    for q in lq:
        topics[q.tag].append(q)
    for topic in topics:
        nq = len(topics[topic])
        total_points = 0
        i = 0
        points = POINTS_IN_TOPIC // nq
        while i < nq:
            if i!= nq - 1:
                topics[topic][i].points = points
                total_points += points
            else:
                topics[topic][i].points = POINTS_IN_TOPIC - total_points
            i +=1


list_questions = []


for file in get_list_file():
    doc = Document(file)
    parse_doc(doc,file.stem.lower(),list_questions)
distribute_points(list_questions)
export_to_xml(list_questions,'text.xml')
print(f'Выгруженно {len(list_questions)} вопросов')
